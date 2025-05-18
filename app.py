import os
import uuid
from flask import Flask, request, jsonify, render_template
import requests
import json
from docx import Document
from docx.shared import Inches
import datetime

app = Flask(__name__)

# Configuration
OLLAMA_BASE_URL = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
EXPORT_DIR = os.path.join(os.path.dirname(__file__), 'exports')
# Simple in-memory storage for novels (replace with a database for persistence)
novels_db = {}

# Ensure export directory exists
if not os.path.exists(EXPORT_DIR):
    os.makedirs(EXPORT_DIR)

# --- Helper Functions ---

def get_ollama_models():
    """Fetches the list of models from the Ollama API."""
    try:
        response = requests.get(f"{OLLAMA_BASE_URL}/api/tags")
        response.raise_for_status() # Raise an exception for bad status codes (4xx or 5xx)
        models_data = response.json()
        # Extract just the model names
        model_names = [model['name'] for model in models_data.get('models', [])]
        return model_names
    except requests.exceptions.RequestException as e:
        print(f"Error connecting to Ollama: {e}")
        return None # Indicate failure
    except json.JSONDecodeError:
        print(f"Error decoding JSON response from Ollama: {response.text}")
        return None
    except Exception as e:
        print(f"An unexpected error occurred while fetching models: {e}")
        return None

# --- API Endpoints ---

@app.route('/')
def index():
    """Serves the main HTML page."""
    return render_template('index.html')

@app.route('/api/models', methods=['GET'])
def list_models():
    """API endpoint to get the list of available Ollama models."""
    models = get_ollama_models()
    if models is None:
        return jsonify({"error": "Could not connect to Ollama or fetch models. Is Ollama running and accessible at " + OLLAMA_BASE_URL + "?"}), 500
    return jsonify(models)

@app.route('/api/novels', methods=['GET'])
def get_novels():
    """API endpoint to list all novels."""
    # Return a list of novels with id and title
    novel_list = [{"id": novel_id, "title": data.get("title", "Untitled")} for novel_id, data in novels_db.items()]
    return jsonify(novel_list)

@app.route('/api/novels', methods=['POST'])
def create_novel():
    """API endpoint to create a new novel."""
    data = request.get_json()
    if not data or 'title' not in data:
        return jsonify({"message": "Title is required"}), 400

    novel_id = str(uuid.uuid4()) # Generate a unique ID
    new_novel = {
        "title": data['title'],
        "content": data.get('content', ''), # Initialize with empty content or provided content
        "characters": "",
        "knowledge": "",
        "prompt_library": ""
        # Add other metadata later (e.g., settings, characters)
    }
    novels_db[novel_id] = new_novel
    print(f"Created novel: {novel_id} - {data['title']}")
    print(f"Current DB: {novels_db}")
    return jsonify({"id": novel_id, "title": new_novel["title"]}), 201 # 201 Created status

@app.route('/api/novels/<novel_id>', methods=['GET'])
def get_novel(novel_id):
    """API endpoint to get details of a specific novel."""
    novel = novels_db.get(novel_id)
    if novel:
        return jsonify({
            "id": novel_id,
            "title": novel.get("title", "Untitled"),
            "characters": novel.get("characters", ""),
            "content": novel.get("content", ""),
            "knowledge": novel.get("knowledge", ""),
            "prompt_library": novel.get("prompt_library", "")
            # Add other fields as needed
        })
    else:
        return jsonify({"message": "Novel not found"}), 404

@app.route('/api/novels/<novel_id>', methods=['PUT', 'PATCH'])
def update_novel(novel_id):
    """API endpoint to update a novel's content (or other fields)."""
    if novel_id not in novels_db:
        return jsonify({"message": "Novel not found"}), 404

    data = request.get_json()
    if not data:
        return jsonify({"message": "No update data provided"}), 400

    # Update fields provided in the request
    if 'title' in data:
        novels_db[novel_id]['title'] = data['title']
    if 'content' in data:
        novels_db[novel_id]['content'] = data['content']
    # Add updates for metadata fields
    if 'characters' in data:
        novels_db[novel_id]['characters'] = data['characters']
    if 'knowledge' in data:
        novels_db[novel_id]['knowledge'] = data['knowledge']
    if 'prompt_library' in data:
        novels_db[novel_id]['prompt_library'] = data['prompt_library']
    # Add updates for other fields (settings, characters, etc.) here

    print(f"Updated novel: {novel_id}")
    return jsonify({"message": "Novel updated successfully", "id": novel_id})

@app.route('/api/novels/<novel_id>', methods=['DELETE'])
def delete_novel(novel_id):
    """API endpoint to delete a specific novel."""
    if novel_id in novels_db:
        deleted_title = novels_db[novel_id].get('title', 'Untitled')
        del novels_db[novel_id]
        print(f"Deleted novel: {novel_id} - {deleted_title}")
        print(f"Current DB: {novels_db}")
        return jsonify({"message": f"Novel '{deleted_title}' deleted successfully"})
    else:
        print(f"Attempted to delete non-existent novel: {novel_id}")
        return jsonify({"message": "Novel not found"}), 404

@app.route('/api/generate', methods=['POST'])
def generate_text():
    """API endpoint to generate text using Ollama."""
    data = request.get_json()
    if not data or 'model' not in data or 'prompt' not in data:
        return jsonify({"error": "Missing required fields: model and prompt"}), 400

    model = data['model']
    prompt = data['prompt']
    stream = data.get('stream', False) # Default to non-streaming

    ollama_payload = {
        "model": model,
        "prompt": prompt,
        "stream": stream
        # Add other Ollama parameters here if needed (e.g., temperature)
    }

    try:
        response = requests.post(f"{OLLAMA_BASE_URL}/api/generate", json=ollama_payload)
        response.raise_for_status()

        # Handle potential streaming response later if needed
        # For now, assume non-streaming
        if stream:
             # TODO: Implement streaming response handling
             return jsonify({"error": "Streaming not yet implemented in this backend"}), 501
        else:
            result = response.json()
            # The actual generated text field might vary slightly depending on Ollama version
            # Common fields are 'response' or 'generated_text' (check Ollama API docs)
            generated_text = result.get('response', '') # Defaulting to 'response' based on common usage
            return jsonify({"generated_text": generated_text})

    except requests.exceptions.RequestException as e:
        print(f"Error contacting Ollama for generation: {e}")
        error_message = str(e)
        # Try to get more specific error from response if available
        try:
            error_detail = response.json().get('error', error_message)
            error_message = f"Ollama API error: {error_detail}"
        except: # Handle cases where response is not JSON or doesn't have 'error'
            pass
        return jsonify({"error": f"Failed to generate text via Ollama. {error_message}"}), 500
    except json.JSONDecodeError:
         print(f"Error decoding JSON response from Ollama generation: {response.text}")
         return jsonify({"error": "Invalid response format from Ollama generation."}), 500
    except Exception as e:
        print(f"An unexpected error occurred during generation: {e}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

# --- Export Endpoints (Server-Side) ---

def generate_export_filename(title, extension):
    """Generates a sanitized filename with title and timestamp."""
    # Sanitize title (allow alphanumeric and CJK)
    sanitized_title = ''.join(c for c in title if c.isalnum() or '\u4e00' <= c <= '\u9fff')
    sanitized_title = sanitized_title or "Untitled" # Fallback if title is empty after sanitizing
    timestamp = datetime.datetime.now().strftime("%Y-%m-%dT%H-%M-%S")
    return f"{sanitized_title}_{timestamp}.{extension}"

@app.route('/api/novels/<novel_id>/export/<file_format>', methods=['POST'])
def export_novel_serverside(novel_id, file_format):
    """API endpoint to export a novel to a file on the server."""
    if novel_id not in novels_db:
        return jsonify({"error": "Novel not found"}), 404

    novel_data = novels_db[novel_id]
    title = novel_data.get('title', 'Untitled')
    content = novel_data.get('content', '')
    characters = novel_data.get('characters', '')
    knowledge = novel_data.get('knowledge', '')
    prompt_library = novel_data.get('prompt_library', '')

    try:
        if file_format == 'txt':
            filename = generate_export_filename(title, 'txt')
            filepath = os.path.join(EXPORT_DIR, filename)
            txt_content = f"标题: {title}\n\n"
            if characters: txt_content += f"角色设定:\n{characters}\n\n"
            if knowledge: txt_content += f"关联知识:\n{knowledge}\n\n"
            if prompt_library: txt_content += f"提示词库:\n{prompt_library}\n\n"
            txt_content += f"正文内容:\n{content}"
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(txt_content)

        elif file_format == 'json':
            filename = generate_export_filename(title, 'json')
            filepath = os.path.join(EXPORT_DIR, filename)
            export_data = {
                'title': title,
                'characters': characters,
                'content': content,
                'knowledge': knowledge,
                'prompt_library': prompt_library
            }
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, ensure_ascii=False, indent=4)

        elif file_format == 'docx':
            filename = generate_export_filename(title, 'docx')
            filepath = os.path.join(EXPORT_DIR, filename)

            document = Document()
            document.add_heading(title, level=0) # Title
            document.add_paragraph()

            if characters:
                document.add_heading('角色设定', level=2)
                for line in characters.split('\n'):
                    document.add_paragraph(line)
                document.add_paragraph()
            if knowledge:
                document.add_heading('关联知识', level=2)
                for line in knowledge.split('\n'):
                    document.add_paragraph(line)
                document.add_paragraph()
            if prompt_library:
                document.add_heading('提示词库', level=2)
                for line in prompt_library.split('\n'):
                    document.add_paragraph(line)
                document.add_paragraph()

            document.add_heading('正文内容', level=1)
            for line in content.split('\n'): # Add content line by line
                document.add_paragraph(line)

            document.save(filepath)

        else:
            return jsonify({"error": "Invalid file format requested"}), 400

        print(f"Exported novel '{title}' ({novel_id}) to {filepath}")
        # Return path relative to server script for confirmation
        relative_path = os.path.relpath(filepath, os.path.dirname(__file__))
        return jsonify({"message": f"文件已成功导出至服务器: {relative_path}", "filepath": relative_path})

    except Exception as e:
        print(f"Error exporting novel {novel_id} to {file_format}: {e}")
        return jsonify({"error": f"导出文件时发生错误: {str(e)}"}), 500

if __name__ == '__main__':
    # Make sure the static folder is configured correctly
    app.run(debug=True) # debug=True enables auto-reloading and detailed errors 